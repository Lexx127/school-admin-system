import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

doc = docx.Document(r"c:\Users\lexxa\school-admin-system\agape_planning_report_v2.docx")
for para in doc.paragraphs:
    if para.text.strip(): print(para.text)

for table in doc.tables:
    for row in table.rows:
        print(" | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]))
